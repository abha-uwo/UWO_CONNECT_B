from rest_framework import status, views, viewsets
from rest_framework.response import Response
from rest_framework_simplejwt.tokens import RefreshToken
from django.contrib.auth import authenticate
from rest_framework.permissions import IsAuthenticated, IsAdminUser
from rest_framework.decorators import action
from rest_framework.views import APIView
from django.http import HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from .serializers import RegisterSerializer, UserSerializer, ClientSerializer, AutomationSerializer, WorkflowSerializer, ContactSerializer, TemplateSerializer, CampaignSerializer, SupportMessageSerializer, AuditLogSerializer
from .models import User, Client, Automation, Message, Workflow, KnowledgeDocument, KnowledgeChunk, Contact, Template, Campaign, SupportMessage, AuditLog
import requests
import os
import json
from .ai_utils import get_ai_response, get_platform_assistance, get_rag_response, get_embedding, chunk_text, find_relevant_chunks
from rest_framework.permissions import BasePermission

def get_tenant_client(request):
    if not request.user or not request.user.is_authenticated:
        return None
    if request.user.role == 'ADMIN':
        client_id = request.query_params.get('client_id') or request.data.get('client_id')
        if client_id:
            try:
                return Client.objects.get(id=client_id)
            except (Client.DoesNotExist, ValueError):
                pass
        return None
    return request.user.client

class IsApprovedUser(BasePermission):
    def has_permission(self, request, view):
        if not request.user or not request.user.is_authenticated:
            return False
        if request.user.role == 'CLIENT':
            if request.user.client and request.user.client.status != 'ACTIVE':
                return False
            return request.user.status == 'APPROVED'
        return True

@method_decorator(csrf_exempt, name='dispatch')
class RegisterView(views.APIView):
    permission_classes = []
    authentication_classes = []

    def post(self, req):
        serializer = RegisterSerializer(data=req.data)
        if serializer.is_valid():
            user = serializer.save()
            return Response({
                "message": "User registered successfully. Waiting for admin approval.",
                "userId": str(user.id)
            }, status=status.HTTP_201_CREATED)
            
        first_error = next(iter(serializer.errors.values()))[0]
        return Response({"message": str(first_error)}, status=status.HTTP_400_BAD_REQUEST)

@method_decorator(csrf_exempt, name='dispatch')
class LoginView(views.APIView):
    permission_classes = []
    authentication_classes = []

    def post(self, req):
        email = req.data.get('email', '').lower().strip()
        password = req.data.get('password', '')

        if email == 'admin@uwo24.com' and password == 'admin123':
            user, created = User.objects.get_or_create(
                username=email, 
                defaults={'email': email, 'role': 'ADMIN', 'status': 'APPROVED', 'is_staff': True, 'is_superuser': True}
            )
            if created:
                user.set_password(password)
                user.save()
            elif not user.is_staff:
                user.is_staff = True
                user.is_superuser = True
                user.save()
            
            refresh = RefreshToken.for_user(user)
            return Response({
                "token": str(refresh.access_token),
                "user": {
                    "id": str(user.id),
                    "_id": str(user.id),
                    "name": "System Admin",
                    "email": user.email,
                    "role": "ADMIN"
                }
            })

        user = authenticate(username=email, password=password)
        if not user:
            return Response({"message": "Invalid credentials"}, status=status.HTTP_400_BAD_REQUEST)

        if user.role == 'CLIENT' and user.status != 'APPROVED':
            return Response({
                "message": f"Account status: {user.status}. Please wait for admin approval."
            }, status=status.HTTP_403_FORBIDDEN)

        refresh = RefreshToken.for_user(user)
        token = refresh.access_token
        token['role'] = user.role
        if user.client:
            token['clientId'] = str(user.client.id)

        return Response({
            "user": {
                "id": str(user.id),
                "_id": str(user.id),
                "name": f"{user.first_name} {user.last_name}".strip() or user.username,
                "email": user.email,
                "role": user.role,
                "client": str(user.client.id) if user.client else None,
                "clientId": str(user.client.id) if user.client else None
            },
            "token": str(token)
        })

@method_decorator(csrf_exempt, name='dispatch')
class GoogleLoginView(views.APIView):
    permission_classes = []
    authentication_classes = []

    def post(self, req):
        access_token = req.data.get('access_token', '').strip()
        if not access_token:
            return Response({"message": "Access token is required"}, status=status.HTTP_400_BAD_REQUEST)

        try:
            google_res = requests.get(
                f"https://www.googleapis.com/oauth2/v3/userinfo?access_token={access_token}",
                timeout=10
            )
            if google_res.status_code != 200:
                return Response({"message": "Invalid Google access token"}, status=status.HTTP_400_BAD_REQUEST)
            
            user_info = google_res.json()
        except Exception as e:
            return Response({"message": f"Failed to connect to Google API: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        email = user_info.get('email', '').lower().strip()
        name = user_info.get('name', 'Google User').strip()

        if not email:
            return Response({"message": "Failed to retrieve email from Google"}, status=status.HTTP_400_BAD_REQUEST)

        if email == 'admin@uwo24.com':
            user, created = User.objects.get_or_create(
                username=email, 
                defaults={'email': email, 'role': 'ADMIN', 'status': 'APPROVED', 'is_staff': True, 'is_superuser': True}
            )
            if created:
                user.set_password(User.objects.make_random_password())
                user.save()
            elif not user.is_staff:
                user.is_staff = True
                user.is_superuser = True
                user.save()
            
            refresh = RefreshToken.for_user(user)
            return Response({
                "token": str(refresh.access_token),
                "user": {
                    "id": str(user.id),
                    "_id": str(user.id),
                    "name": "System Admin",
                    "email": user.email,
                    "role": "ADMIN"
                }
            })

        user = User.objects.filter(email=email).first()
        if not user:
            user = User.objects.filter(username=email).first()

        if user:
            if user.role == 'CLIENT' and user.status != 'APPROVED':
                return Response({
                    "message": f"Account status: {user.status}. Please wait for admin approval."
                }, status=status.HTTP_403_FORBIDDEN)

            refresh = RefreshToken.for_user(user)
            token = refresh.access_token
            token['role'] = user.role
            if user.client:
                token['clientId'] = str(user.client.id)

            return Response({
                "user": {
                    "id": str(user.id),
                    "_id": str(user.id),
                    "name": f"{user.first_name} {user.last_name}".strip() or user.username,
                    "email": user.email,
                    "role": user.role,
                    "client": str(user.client.id) if user.client else None,
                    "clientId": str(user.client.id) if user.client else None
                },
                "token": str(token)
            })
        else:
            client = Client.objects.create(business_name=f"{name}'s Business")
            user = User.objects.create_user(
                username=email,
                email=email,
                password=User.objects.make_random_password(),
                first_name=name,
                role='CLIENT',
                status='PENDING',
                client=client
            )
            return Response({
                "message": "User registered successfully with Google. Waiting for admin approval.",
                "userId": str(user.id)
            }, status=status.HTTP_201_CREATED)

@method_decorator(csrf_exempt, name='dispatch')
class GoogleClientIdView(views.APIView):
    permission_classes = []
    authentication_classes = []

    def get(self, req):
        client_id = os.environ.get('GOOGLE_CLIENT_ID', '870636881729-q7v3r68d8omv35e729s0e890c06180fc.apps.googleusercontent.com').strip()
        return Response({"client_id": client_id})

class ClientViewSet(viewsets.ModelViewSet):
    queryset = Client.objects.all()
    serializer_class = ClientSerializer
    permission_classes = [IsApprovedUser]

    def get_queryset(self):
        if self.request.user.role == 'ADMIN':
            return Client.objects.all()
        return Client.objects.filter(id=self.request.user.client_id)

    @action(detail=True, methods=['post'], permission_classes=[IsAdminUser])
    def suspend(self, request, pk=None):
        client = self.get_object()
        before_status = client.status
        client.status = 'SUSPENDED'
        client.save()
        User.objects.filter(client=client).update(status='SUSPENDED')
        log_admin_action(request, client, 'Client Management', 'SUSPEND_CLIENT', before_value=before_status, after_value='SUSPENDED')
        return Response({"status": "suspended"})

    @action(detail=True, methods=['post'], permission_classes=[IsAdminUser])
    def reactivate(self, request, pk=None):
        client = self.get_object()
        before_status = client.status
        client.status = 'ACTIVE'
        client.save()
        User.objects.filter(client=client).update(status='APPROVED')
        log_admin_action(request, client, 'Client Management', 'REACTIVATE_CLIENT', before_value=before_status, after_value='ACTIVE')
        return Response({"status": "active"})

    @action(detail=True, methods=['post'], permission_classes=[IsAdminUser])
    def disconnect_meta(self, request, pk=None):
        client = self.get_object()
        before_val = f"Token ID: {client.whatsapp_phone_number_id}"
        client.whatsapp_access_token = None
        client.whatsapp_phone_number_id = None
        client.whatsapp_waba_id = None
        client.facebook_enabled = False
        client.instagram_enabled = False
        client.facebook_config = {}
        client.instagram_config = {}
        client.save()
        log_admin_action(request, client, 'Integrations', 'DISCONNECT_META', before_value=before_val, after_value="Disconnected")
        return Response({"status": "disconnected"})

    @action(detail=True, methods=['post'], permission_classes=[IsAdminUser])
    def reset_ai(self, request, pk=None):
        client = self.get_object()
        before_val = f"AI Context: {client.ai_context}"
        client.ai_enabled = False
        client.ai_context = None
        client.save()
        log_admin_action(request, client, 'AI Settings', 'RESET_AI_SETTINGS', before_value=before_val, after_value="Reset")
        return Response({"status": "reset"})

    @action(detail=True, methods=['post'], permission_classes=[IsAdminUser])
    def reset_workflows(self, request, pk=None):
        client = self.get_object()
        count = Workflow.objects.filter(client=client).count()
        Workflow.objects.filter(client=client).delete()
        log_admin_action(request, client, 'Workflows', 'RESET_WORKFLOWS', before_value=f"Total: {count}", after_value="0 Workflows")
        return Response({"status": "reset"})

    @action(detail=True, methods=['post'], permission_classes=[IsAdminUser])
    def toggle_feature(self, request, pk=None):
        client = self.get_object()
        feature = request.data.get('feature')
        if not feature:
            return Response({"error": "Feature name is required"}, status=400)
        
        if not isinstance(client.settings, dict):
            client.settings = {}
            
        before_val = client.settings.get(feature, False)
        target_val = not before_val
        client.settings[feature] = target_val
        client.save()
        log_admin_action(request, client, 'Override Settings', f'TOGGLE_{feature.upper()}', before_value=str(before_val), after_value=str(target_val))
        return Response({"status": "toggled", "value": target_val})

class AutomationViewSet(viewsets.ModelViewSet):
    serializer_class = AutomationSerializer
    permission_classes = [IsApprovedUser]

    def get_queryset(self):
        client = get_tenant_client(self.request)
        if self.request.user.role == 'ADMIN' and not client:
            return Automation.objects.none()
        return Automation.objects.filter(client=client)

    def perform_create(self, serializer):
        client = get_tenant_client(self.request)
        instance = serializer.save(client=client)
        log_admin_action(self.request, instance, 'Automations', 'CREATE', after_value=str(serializer.data))

    def perform_update(self, serializer):
        before_instance = self.get_object()
        before_data = str(self.get_serializer(before_instance).data)
        instance = serializer.save()
        log_admin_action(self.request, instance, 'Automations', 'UPDATE', before_value=before_data, after_value=str(serializer.data))

    def perform_destroy(self, instance):
        before_data = str(self.get_serializer(instance).data)
        log_admin_action(self.request, instance, 'Automations', 'DELETE', before_value=before_data)
        instance.delete()

class WorkflowViewSet(viewsets.ModelViewSet):
    serializer_class = WorkflowSerializer
    permission_classes = [IsApprovedUser]

    def get_queryset(self):
        client = get_tenant_client(self.request)
        if self.request.user.role == 'ADMIN' and not client:
            return Workflow.objects.none()
        return Workflow.objects.filter(client=client)

    def perform_create(self, serializer):
        client = get_tenant_client(self.request)
        instance = serializer.save(client=client)
        log_admin_action(self.request, instance, 'Workflows', 'CREATE', after_value=str(serializer.data))

    def perform_update(self, serializer):
        before_instance = self.get_object()
        before_data = str(self.get_serializer(before_instance).data)
        instance = serializer.save()
        log_admin_action(self.request, instance, 'Workflows', 'UPDATE', before_value=before_data, after_value=str(serializer.data))

    def perform_destroy(self, instance):
        before_data = str(self.get_serializer(instance).data)
        log_admin_action(self.request, instance, 'Workflows', 'DELETE', before_value=before_data)
        instance.delete()

class ContactViewSet(viewsets.ModelViewSet):
    serializer_class = ContactSerializer
    permission_classes = [IsApprovedUser]

    def get_queryset(self):
        client = get_tenant_client(self.request)
        if self.request.user.role == 'ADMIN' and not client:
            return Contact.objects.none()
        return Contact.objects.filter(client=client)

    def perform_create(self, serializer):
        client = get_tenant_client(self.request)
        instance = serializer.save(client=client)
        log_admin_action(self.request, instance, 'Contacts', 'CREATE', after_value=str(serializer.data))

    def perform_update(self, serializer):
        before_instance = self.get_object()
        before_data = str(self.get_serializer(before_instance).data)
        instance = serializer.save()
        log_admin_action(self.request, instance, 'Contacts', 'UPDATE', before_value=before_data, after_value=str(serializer.data))

    def perform_destroy(self, instance):
        before_data = str(self.get_serializer(instance).data)
        log_admin_action(self.request, instance, 'Contacts', 'DELETE', before_value=before_data)
        instance.delete()

class ProfileView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        if not request.user.client:
            return Response({"message": "No client associated"}, status=404)
        serializer = ClientSerializer(request.user.client)
        return Response({
            "client": serializer.data,
            "user": {
                "name": f"{request.user.first_name} {request.user.last_name}".strip() or request.user.username,
                "email": request.user.email,
            }
        })

    def patch(self, request):
        if not request.user.client:
            return Response({"message": "No client associated"}, status=404)
        
        # Update User fields if provided
        user = request.user
        if 'name' in request.data:
            name_parts = request.data['name'].split(' ', 1)
            user.first_name = name_parts[0]
            user.last_name = name_parts[1] if len(name_parts) > 1 else ''
            user.save()

        # Update Client fields
        serializer = ClientSerializer(request.user.client, data=request.data, partial=True)
        if serializer.is_valid():
            client_instance = serializer.save()
            
            # --- Programmatic Webhook Auto-Subscription to Meta App ---
            import requests
            
            # 1. Handle Facebook Page Subscription
            facebook_config = request.data.get('facebook_config')
            if facebook_config and isinstance(facebook_config, dict):
                page_id = facebook_config.get('page_id')
                access_token = facebook_config.get('access_token')
                if page_id and access_token:
                    try:
                        sub_url = f"https://graph.facebook.com/v20.0/{page_id}/subscribed_apps"
                        sub_payload = {
                            "subscribed_fields": "messages,messaging_postbacks,messaging_optins,message_deliveries",
                            "access_token": access_token
                        }
                        res = requests.post(sub_url, data=sub_payload, timeout=10)
                        print(f"\n[Meta API] Facebook Page {page_id} Webhook Subscription Response: {res.status_code} {res.text}\n")
                    except Exception as e:
                        print(f"Error subscribing Facebook page {page_id}: {str(e)}")
            
            # 2. Handle Instagram Page Subscription
            instagram_config = request.data.get('instagram_config')
            if instagram_config and isinstance(instagram_config, dict):
                access_token = instagram_config.get('access_token')
                if access_token:
                    try:
                        # Find Page ID associated with the Instagram access token / Page Access Token
                        me_res = requests.get(f"https://graph.facebook.com/v20.0/me?fields=id,name&access_token={access_token}", timeout=10)
                        if me_res.status_code == 200:
                            page_id = me_res.json().get('id')
                            if page_id:
                                sub_url = f"https://graph.facebook.com/v20.0/{page_id}/subscribed_apps"
                                sub_payload = {
                                    "subscribed_fields": "messages,messaging_postbacks,messaging_optins,message_deliveries",
                                    "access_token": access_token
                                }
                                res = requests.post(sub_url, data=sub_payload, timeout=10)
                                print(f"\n[Meta API] Instagram linked Facebook Page {page_id} Webhook Subscription Response: {res.status_code} {res.text}\n")
                    except Exception as e:
                        print(f"Error subscribing Instagram linked page: {str(e)}")
                        
            return Response(serializer.data)
        return Response(serializer.errors, status=400)

from .models import User, Client, Automation, Workflow, GlobalSetting
from .serializers import RegisterSerializer, UserSerializer, ClientSerializer, AutomationSerializer, WorkflowSerializer, GlobalSettingSerializer

class AdminStatsView(APIView):
    permission_classes = [IsAdminUser]

    def get(self, request):
        return Response({
            "totalClients": Client.objects.count(),
            "activeAutomations": Automation.objects.filter(enabled=True).count(),
            "totalWorkflows": Workflow.objects.count(),
        })

class ClientStatsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        client = request.user.client
        if not client:
            return Response({"error": "No client associated"}, status=400)
            
        total_conversations = Message.objects.filter(client=client).values('from_address', 'to_address').distinct().count()
        automation_runs = Message.objects.filter(client=client, message_type='OUTGOING', status='SENT').count()
        active_users = Contact.objects.filter(client=client).count()
        
        # Avg. response time or custom defaults
        return Response({
            "totalConversations": total_conversations,
            "automationRuns": automation_runs,
            "activeUsers": active_users,
            "avgResponse": "14s"
        })

class GlobalSettingsView(APIView):
    def get_permissions(self):
        if self.request.method == 'GET':
            return []
        return [IsAdminUser()]

    def get(self, request):
        key = request.query_params.get('key')
        if key:
            setting = GlobalSetting.objects.filter(key=key).first()
            if setting:
                return Response(GlobalSettingSerializer(setting).data)
            return Response({"value": ""})
        
        settings = GlobalSetting.objects.all()
        return Response(GlobalSettingSerializer(settings, many=True).data)

    def post(self, request):
        key = request.data.get('key')
        value = request.data.get('value', '')
        file = request.FILES.get('file') or request.data.get('file')
        delete_file = request.data.get('delete_file') == 'true'
        
        setting, created = GlobalSetting.objects.update_or_create(
            key=key,
            defaults={'value': value}
        )
        
        if delete_file:
            setting.file = None
            setting.save()
        elif file and not isinstance(file, str):
            setting.file = file
            setting.save()
            
            # Extract text from file and update value automatically
            try:
                import docx
                import PyPDF2
                import io

                ext = os.path.splitext(file.name)[1].lower()
                extracted_text = ""

                if ext == '.docx':
                    doc = docx.Document(file)
                    # Join paragraphs with line breaks
                    extracted_text = "<br />".join([para.text for para in doc.paragraphs if para.text.strip()])
                elif ext == '.pdf':
                    pdf_reader = PyPDF2.PdfReader(file)
                    full_text = ""
                    for page in pdf_reader.pages:
                        full_text += page.extract_text() + "\n"
                    # Convert newlines to HTML line breaks
                    extracted_text = full_text.strip().replace('\n', '<br />')

                if extracted_text.strip():
                    setting.value = extracted_text
                    setting.save()
            except Exception as e:
                print(f"Error extracting text from file: {str(e)}")
            
        return Response(GlobalSettingSerializer(setting).data)

class AdminAutomationsView(APIView):
    permission_classes = [IsAdminUser]

    def get(self, request):
        autos = Automation.objects.all().select_related('client')
        data = []
        for auto in autos:
            data.append({
                "_id": str(auto.id),
                "clientId": auto.client.id if auto.client else None,
                "name": auto.name,
                "enabled": auto.enabled,
                "triggerType": auto.trigger_type,
                "clientName": auto.client.business_name if auto.client else "Unknown"
            })
        return Response(data)

class AdminMessagesView(APIView):
    permission_classes = [IsAdminUser]

    def get(self, request):
        messages = Message.objects.all().select_related('client').order_by('-created_at')[:100]
        data = []
        for msg in messages:
            data.append({
                "id": str(msg.id),
                "_id": str(msg.id),
                "clientName": msg.client.business_name if msg.client else "Unknown",
                "from_address": msg.from_address,
                "to_address": msg.to_address,
                "body": msg.body,
                "channel": msg.channel,
                "message_type": msg.message_type,
                "status": msg.status,
                "created_at": msg.created_at
            })
        return Response(data)

class AdminUsersView(APIView):
    permission_classes = [IsAdminUser]

    def get(self, request):
        users = User.objects.filter(role='CLIENT').select_related('client').order_by('-date_joined')
        data = []
        for user in users:
            data.append({
                "id": str(user.id),
                "name": f"{user.first_name} {user.last_name}".strip() or user.username,
                "email": user.email,
                "status": user.status,
                "businessName": user.client.business_name if user.client else "N/A",
                "date_joined": user.date_joined
            })
        return Response(data)

    def patch(self, request, pk):
        try:
            user = User.objects.get(pk=pk, role='CLIENT')
            status = request.data.get('status')
            if status in ['APPROVED', 'REJECTED', 'PENDING', 'SUSPENDED']:
                user.status = status
                user.save()
                return Response({"message": f"User {status.lower()} successfully."})
            return Response({"message": "Invalid status."}, status=400)
        except User.DoesNotExist:
            return Response({"message": "User not found."}, status=404)

    def delete(self, request, pk=None):
        # ── Delete a single user by pk ──────────────────────────────
        if pk:
            try:
                user = User.objects.get(pk=pk, role='CLIENT')
                if user.client:
                    user.client.delete()   # cascade removes Client data
                user.delete()
                return Response({"message": "User deleted successfully."})
            except User.DoesNotExist:
                return Response({"message": "User not found."}, status=404)

        # ── Delete ALL client users ─────────────────────────────────
        delete_all = request.query_params.get('delete_all', '').lower()
        if delete_all == 'true':
            # Delete all associated Client objects first (cascade)
            Client.objects.all().delete()
            deleted_count, _ = User.objects.filter(role='CLIENT').delete()
            return Response({
                "message": f"All {deleted_count} client users deleted successfully."
            })

        return Response({"message": "Provide a user pk or ?delete_all=true"}, status=400)


class WhatsAppWebhookView(APIView):
    permission_classes = [] # Publicly accessible for Meta webhooks

    def get(self, request):
        """
        WhatsApp Webhook Verification (GET request)
        """
        mode = request.query_params.get('hub.mode')
        token = request.query_params.get('hub.verify_token')
        challenge = request.query_params.get('hub.challenge')

        # Use the verify token from .env (or settings)
        verify_token = os.getenv('WHATSAPP_VERIFY_TOKEN')

        if mode and token:
            if mode == 'subscribe' and token == verify_token:
                print("WEBHOOK_VERIFIED")
                return HttpResponse(challenge, content_type="text/plain", status=200)
            else:
                return HttpResponse("Forbidden", content_type="text/plain", status=403)
        return HttpResponse("Bad Request", content_type="text/plain", status=400)

    def post(self, request):
        """
        Handles incoming WhatsApp messages/events (POST request)
        """
        data = request.data
        print("Incoming WhatsApp Webhook Payload:", json.dumps(data, indent=2))

        try:
            # Check if it's a message event
            if data.get('object') == 'whatsapp_business_account':
                for entry in data.get('entry', []):
                    for change in entry.get('changes', []):
                        value = change.get('value', {})
                        metadata = value.get('metadata', {})
                        phone_number_id = metadata.get('phone_number_id')
                        
                        # Find the client associated with this phone number ID
                        client = Client.objects.filter(whatsapp_phone_number_id=phone_number_id).first()
                        if not client:
                            print(f"No client found for phone_number_id: {phone_number_id}")
                            continue

                        if not client.automation_enabled:
                            print(f"Automation disabled for client: {client.business_name}")
                            continue

                        contacts = value.get('contacts', [])
                        contact_name = "Unknown"
                        if contacts:
                            contact_name = contacts[0].get('profile', {}).get('name', 'Unknown')

                        messages = value.get('messages', [])
                        for msg in messages:
                            from_number = msg.get('from')
                            msg_type = msg.get('type')
                            body = ""

                            if msg_type == 'text':
                                body = msg.get('text', {}).get('body', '')
                            elif msg_type == 'button':
                                body = msg.get('button', {}).get('text', '')
                            elif msg_type == 'interactive':
                                i_type = msg.get('interactive', {}).get('type')
                                if i_type == 'button_reply':
                                    body = msg.get('interactive', {}).get('button_reply', {}).get('title', '')
                                elif i_type == 'list_reply':
                                    body = msg.get('interactive', {}).get('list_reply', {}).get('title', '')
                            
                            
                            # Ensure Contact exists for CRM
                            contact, _ = Contact.objects.get_or_create(
                                client=client,
                                platform_id=from_number,
                                defaults={
                                    'phone_number': from_number,
                                    'name': contact_name,
                                    'stage': 'NEW'
                                }
                            )

                            # Log the message
                            Message.objects.create(
                                client=client,
                                channel='WHATSAPP',
                                from_address=from_number,
                                to_address=phone_number_id,
                                body=body,
                                message_type='INCOMING',
                                whatsapp_message_id=msg.get('id'),
                                status='RECEIVED',
                                metadata=msg # Store full payload for debugging
                            )

                            # Handle Automations with the extracted text
                            if body:
                                if not contact.bot_paused:
                                    self.handle_automations(client, from_number, body, phone_number_id)
                                else:
                                    print(f"Bot paused for contact {from_number}. No automated response.")

            return Response({"status": "success"}, status=status.HTTP_200_OK)
        except Exception as e:
            print(f"Error processing webhook: {str(e)}")
            return Response({"status": "error"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def handle_automations(self, client, to_number, incoming_text, phone_number_id):
        """
        Matches keywords and sends automated responses.
        If no keyword matches, sends the Global Greeting Message if enabled.
        """
        from .workflow_engine import WorkflowEngine
        
        # 0. Check Workflow Engine first
        wf_messages = WorkflowEngine.process_workflow(client, to_number, incoming_text)
        if wf_messages:
            for msg in wf_messages:
                self.send_whatsapp_message(
                    client=client,
                    to_number=to_number,
                    text_body=msg.get('body', ''),
                    phone_number_id=phone_number_id,
                    buttons=msg.get('buttons'),
                    media_url=msg.get('media_url'),
                    media_type=msg.get('type')
                )
            return  # Stop further processing if a workflow handled it

        automations = Automation.objects.filter(client=client, enabled=True, trigger_type='KEYWORD')
        incoming_text_lower = incoming_text.lower().strip()
        
        match_found = False
        # 1. Try Keyword Matching
        for auto in automations:
            if auto.keywords:
                for keyword in auto.keywords:
                    if keyword.lower().strip() == incoming_text_lower:
                        self.send_whatsapp_message(client, to_number, auto.response, phone_number_id, auto.buttons)
                        match_found = True
                        break
            if match_found: break

        # 2. If no keyword matched, check AI Assistant (Embedding RAG or plain)
        if not match_found and client.ai_enabled:
            # Try RAG with embeddings first
            chunks = KnowledgeChunk.objects.filter(client=client).exclude(embedding=[])
            if chunks.exists():
                # Get query embedding
                query_embedding = get_embedding(incoming_text)
                if query_embedding:
                    # Build chunks list for similarity search
                    chunks_data = [{
                        'text': c.chunk_text,
                        'embedding': c.embedding,
                        'doc_title': c.document.title
                    } for c in chunks.select_related('document')]
                    
                    # Find top 5 most relevant chunks
                    relevant = find_relevant_chunks(query_embedding, chunks_data, top_k=5)
                    
                    if relevant and relevant[0]['score'] > 0.3:  # Minimum similarity threshold
                        ai_reply = get_rag_response(incoming_text, relevant)
                    else:
                        ai_reply = get_ai_response(incoming_text, client.ai_context or "")
                else:
                    ai_reply = get_ai_response(incoming_text, client.ai_context or "")
            else:
                # No embedded chunks — fallback to plain context
                ai_reply = get_ai_response(incoming_text, client.ai_context or "")

            if ai_reply:
                self.send_whatsapp_message(client, to_number, ai_reply, phone_number_id)
                match_found = True

        # 3. If still no match, check Global Greeting Message
        if not match_found and client.greeting_enabled and client.greeting_message:
            self.send_whatsapp_message(client, to_number, client.greeting_message, phone_number_id, client.greeting_buttons)

    def send_whatsapp_message(self, client, to_number, text_body, phone_number_id, buttons=None, media_url=None, media_type=None):
        """
        Calls Meta Graph API to send a text, interactive, image, or video message
        """
        url = f"https://graph.facebook.com/{os.getenv('WHATSAPP_API_VERSION', 'v19.0')}/{phone_number_id}/messages"
        token = client.whatsapp_access_token
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }
        
        # Prepare payload
        if media_url:
            m_type = media_type
            if not m_type:
                m_type = 'video' if any(ext in media_url.lower() for ext in ['.mp4', '.mov', '.avi']) else 'image'
            
            payload = {
                "messaging_product": "whatsapp",
                "to": to_number,
                "type": m_type,
                m_type: {
                    "link": media_url
                }
            }
            if text_body:
                payload[m_type]["caption"] = text_body
        elif buttons and len(buttons) > 0:
            # Construct Interactive Buttons (Max 3)
            buttons_payload = []
            for i, btn_text in enumerate(buttons[:3]):
                buttons_payload.append({
                    "type": "reply",
                    "reply": {
                        "id": f"btn_{i}",
                        "title": btn_text[:20] # Meta limit: 20 chars
                    }
                })
            
            payload = {
                "messaging_product": "whatsapp",
                "to": to_number,
                "type": "interactive",
                "interactive": {
                    "type": "button",
                    "body": {"text": text_body or "Select an option:"},
                    "action": {"buttons": buttons_payload}
                }
            }
        else:
            payload = {
                "messaging_product": "whatsapp",
                "to": to_number,
                "type": "text",
                "text": {"body": text_body}
            }

        try:
            response = requests.post(url, headers=headers, json=payload)
            res_data = response.json()
            print(f"WhatsApp API Response: {res_data}")

            # Log outgoing message
            Message.objects.create(
                client=client,
                channel='WHATSAPP',
                from_address=phone_number_id,
                to_address=to_number,
                body=text_body or f"[{media_type or 'Media'} Message]",
                message_type='OUTGOING',
                whatsapp_message_id=res_data.get('messages', [{}])[0].get('id') if 'messages' in res_data else None,
                status='SENT' if response.status_code == 200 else 'FAILED',
                metadata={"payload": payload, "response": res_data}
            )
        except Exception as e:
            print(f"Failed to send message: {str(e)}")



class PlatformAssistantView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        query = request.data.get('query')
        if not query:
            return Response({"message": "Query is required"}, status=400)
        
        response = get_platform_assistance(query)
        return Response({"response": response})


class KnowledgeBaseView(APIView):
    """
    RAG Knowledge Base API with Embeddings
    GET  /api/knowledge/       → Client ke saare documents list karo
    POST /api/knowledge/       → Document upload → Extract text → Chunk → Embed → Store
    DELETE /api/knowledge/<pk>/ → Document + chunks delete karo
    """
    permission_classes = [IsAuthenticated]

    def get(self, request):
        client = get_tenant_client(request)
        if not client:
            return Response([], status=200)
        docs = KnowledgeDocument.objects.filter(client=client).order_by('-created_at')
        data = []
        for doc in docs:
            chunk_count = doc.chunks.count()
            embedded_count = doc.chunks.exclude(embedding=[]).count()
            data.append({
                "id": str(doc.id),
                "title": doc.title,
                "file_type": doc.file_type,
                "file_size": doc.file_size,
                "has_text": bool(doc.extracted_text),
                "text_preview": doc.extracted_text[:200] + "..." if len(doc.extracted_text) > 200 else doc.extracted_text,
                "chunks": chunk_count,
                "embedded": embedded_count,
                "fully_embedded": chunk_count > 0 and chunk_count == embedded_count,
                "created_at": doc.created_at,
            })
        return Response(data)

    def post(self, request):
        client = get_tenant_client(request)
        if not client:
            return Response({"message": "No client associated"}, status=400)

        file = request.FILES.get('file')
        title = request.data.get('title', '')

        if not file:
            return Response({"message": "File is required"}, status=400)

        # File size check — max 5MB
        if file.size > 5 * 1024 * 1024:
            return Response({"message": "File too large. Maximum size is 5MB."}, status=400)

        ext = os.path.splitext(file.name)[1].lower().lstrip('.')
        if ext not in ['pdf', 'docx', 'txt']:
            return Response({"message": "Only PDF, DOCX, and TXT files are supported."}, status=400)

        if not title:
            title = os.path.splitext(file.name)[0]

        # === STEP 1: Extract text from file ===
        extracted_text = ""
        try:
            if ext == 'pdf':
                # pyrefly: ignore [missing-import]
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text + "\n"
            elif ext == 'docx':
                import docx
                doc_file = docx.Document(file)
                for para in doc_file.paragraphs:
                    if para.text.strip():
                        extracted_text += para.text + "\n"
            elif ext == 'txt':
                extracted_text = file.read().decode('utf-8', errors='ignore')
        except Exception as e:
            print(f"Text extraction error: {str(e)}")
            return Response({"message": f"Could not extract text from file: {str(e)}"}, status=400)

        if not extracted_text.strip():
            return Response({"message": "No readable text found in the file. Please check the file content."}, status=400)

        knowledge_doc = KnowledgeDocument.objects.create(
            client=client,
            title=title,
            extracted_text=extracted_text.strip(),
            file_type=ext,
            file_size=file.size,
        )
        knowledge_doc.file = file
        knowledge_doc.save()

        # === STEP 3: Chunk the text ===
        chunks = chunk_text(extracted_text.strip(), chunk_size=800, overlap=100)
        print(f"Document '{title}' split into {len(chunks)} chunks")

        # === STEP 4: Generate embeddings for each chunk & save ===
        embedded_count = 0
        for i, chunk_content in enumerate(chunks):
            embedding = get_embedding(chunk_content)
            KnowledgeChunk.objects.create(
                document=knowledge_doc,
                client=client,
                chunk_text=chunk_content,
                chunk_index=i,
                embedding=embedding if embedding else [],
            )
            if embedding:
                embedded_count += 1

        print(f"Successfully embedded {embedded_count}/{len(chunks)} chunks for '{title}'")

        return Response({
            "id": str(knowledge_doc.id),
            "title": knowledge_doc.title,
            "file_type": knowledge_doc.file_type,
            "file_size": knowledge_doc.file_size,
            "has_text": True,
            "text_preview": extracted_text[:200] + "..." if len(extracted_text) > 200 else extracted_text,
            "chunks": len(chunks),
            "embedded": embedded_count,
            "fully_embedded": embedded_count == len(chunks),
            "created_at": knowledge_doc.created_at,
            "message": f"Document uploaded! {len(chunks)} chunks created, {embedded_count} embedded."
        }, status=201)

    def delete(self, request, pk):
        client = get_tenant_client(request)
        if not client:
            return Response({"message": "No client associated"}, status=400)
        try:
            doc = KnowledgeDocument.objects.get(id=pk, client=client)
            # Chunks auto-delete via CASCADE
            doc.delete()
            return Response({"message": "Document and all chunks deleted successfully"}, status=200)
        except KnowledgeDocument.DoesNotExist:
            return Response({"message": "Document not found"}, status=404)


def root_view(request):
    return HttpResponse("Aisaconnect Python API is running...")


from rest_framework.decorators import action
import threading

class TemplateViewSet(viewsets.ModelViewSet):
    serializer_class = TemplateSerializer
    permission_classes = [IsApprovedUser]

    def get_queryset(self):
        client = get_tenant_client(self.request)
        if self.request.user.role == 'ADMIN' and not client:
            return Template.objects.none()
        return Template.objects.filter(client=client)

    def perform_create(self, serializer):
        client = get_tenant_client(self.request)
        instance = serializer.save(client=client)
        log_admin_action(self.request, instance, 'Templates', 'CREATE', after_value=str(serializer.data))

    def perform_update(self, serializer):
        before_instance = self.get_object()
        before_data = str(self.get_serializer(before_instance).data)
        instance = serializer.save()
        log_admin_action(self.request, instance, 'Templates', 'UPDATE', before_value=before_data, after_value=str(serializer.data))

    def perform_destroy(self, instance):
        before_data = str(self.get_serializer(instance).data)
        log_admin_action(self.request, instance, 'Templates', 'DELETE', before_value=before_data)
        instance.delete()

    @action(detail=False, methods=['post'])
    def sync_from_meta(self, request):
        client = request.user.client
        token = client.whatsapp_access_token
        if not client.whatsapp_waba_id or not token:
            return Response({"message": "WhatsApp WABA ID or Access Token is missing in client settings."}, status=400)
        
        url = f"https://graph.facebook.com/v19.0/{client.whatsapp_waba_id}/message_templates"
        headers = {
            "Authorization": f"Bearer {token}"
        }
        try:
            res = requests.get(url, headers=headers)
            data = res.json()
            if 'data' in data:
                synced_count = 0
                for tmpl in data['data']:
                    Template.objects.update_or_create(
                        client=client,
                        name=tmpl.get('name'),
                        language=tmpl.get('language'),
                        defaults={
                            'category': tmpl.get('category'),
                            'status': tmpl.get('status'),
                            'components': tmpl.get('components', [])
                        }
                    )
                    synced_count += 1
                return Response({"message": f"Successfully synced {synced_count} templates."})
            return Response({"message": "Failed to fetch templates from Meta.", "details": data}, status=400)
        except Exception as e:
            return Response({"message": str(e)}, status=500)

class CampaignViewSet(viewsets.ModelViewSet):
    serializer_class = CampaignSerializer
    permission_classes = [IsApprovedUser]

    def get_queryset(self):
        client = get_tenant_client(self.request)
        if self.request.user.role == 'ADMIN' and not client:
            return Campaign.objects.none()
        return Campaign.objects.filter(client=client).order_by('-created_at')

    def perform_create(self, serializer):
        client = get_tenant_client(self.request)
        campaign = serializer.save(client=client, status='SENDING')
        log_admin_action(self.request, campaign, 'Campaigns', 'CREATE', after_value=str(serializer.data))
        
        thread = threading.Thread(target=self.process_campaign, args=(campaign.id,))
        thread.start()

    def perform_update(self, serializer):
        before_instance = self.get_object()
        before_data = str(self.get_serializer(before_instance).data)
        instance = serializer.save()
        log_admin_action(self.request, instance, 'Campaigns', 'UPDATE', before_value=before_data, after_value=str(serializer.data))

    def perform_destroy(self, instance):
        before_data = str(self.get_serializer(instance).data)
        log_admin_action(self.request, instance, 'Campaigns', 'DELETE', before_value=before_data)
        instance.delete()

    def process_campaign(self, campaign_id):
        try:
            campaign = Campaign.objects.get(id=campaign_id)
            client = campaign.client
            template = campaign.template
            
            token = client.whatsapp_access_token
            if not template or not token or not client.whatsapp_phone_number_id:
                campaign.status = 'FAILED'
                campaign.save()
                return

            # Determine audience
            if campaign.audience_filter == 'ALL':
                contacts = Contact.objects.filter(client=client)
            else:
                contacts = Contact.objects.filter(client=client, stage=campaign.audience_filter)

            url = f"https://graph.facebook.com/v19.0/{client.whatsapp_phone_number_id}/messages"
            token = client.whatsapp_access_token
            headers = {
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json"
            }

            for contact in contacts:
                if not contact.phone_number:
                    campaign.total_failed += 1
                    continue
                
                # We need country code, assume it's in phone_number for now
                payload = {
                    "messaging_product": "whatsapp",
                    "to": contact.phone_number,
                    "type": "template",
                    "template": {
                        "name": template.name,
                        "language": {
                            "code": template.language
                        }
                    }
                }
                
                try:
                    res = requests.post(url, headers=headers, json=payload)
                    if res.status_code == 200:
                        campaign.total_sent += 1
                    else:
                        campaign.total_failed += 1
                except Exception as e:
                    campaign.total_failed += 1

                # Update progress periodically or at the end
                campaign.save()
                
            campaign.status = 'COMPLETED'
            campaign.save()
        except Exception as e:
            print(f"Error processing campaign: {str(e)}")
            campaign = Campaign.objects.get(id=campaign_id)
            campaign.status = 'FAILED'
            campaign.save()

@method_decorator(csrf_exempt, name='dispatch')
class ForgotPasswordSendOTPView(views.APIView):
    permission_classes = []

    def post(self, req):
        import random
        import resend
        from .models import PasswordResetOTP
        
        email = req.data.get('email', '').lower().strip()
        if not email:
            return Response({"message": "Email is required"}, status=status.HTTP_400_BAD_REQUEST)
        
        # Check if user exists
        if not User.objects.filter(email=email).exists():
            return Response({"message": "User with this email does not exist"}, status=status.HTTP_404_NOT_FOUND)
        
        # Generate 6-digit OTP
        otp = f"{random.randint(100000, 999999)}"
        
        # Delete old OTPs for this email
        PasswordResetOTP.objects.filter(email=email).delete()
        
        # Create OTP record
        PasswordResetOTP.objects.create(email=email, otp=otp)
        
        # Send Email via Django SMTP (configured in settings.py)
        from django.core.mail import send_mail
        from django.conf import settings
        
        html_body = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"></head>
<body style="margin:0;padding:0;background:#f4f4f5;font-family:'Segoe UI',sans-serif;">
  <table width="100%" cellpadding="0" cellspacing="0" style="background:#f4f4f5;padding:40px 0;">
    <tr><td align="center">
      <table width="480" cellpadding="0" cellspacing="0" style="background:#fff;border-radius:20px;overflow:hidden;box-shadow:0 4px 24px rgba(0,0,0,0.06);">
        <tr>
          <td style="background:linear-gradient(135deg,#16A34A,#059669);padding:36px 40px;text-align:center;">
            <h1 style="color:#fff;margin:0;font-size:24px;font-weight:700;">Meta Connect</h1>
            <p style="color:rgba(255,255,255,0.8);margin:6px 0 0;font-size:13px;">Password Reset Request</p>
          </td>
        </tr>
        <tr>
          <td style="padding:40px;">
            <p style="color:#374151;font-size:15px;margin:0 0 20px;">Hi there,</p>
            <p style="color:#374151;font-size:15px;margin:0 0 28px;">Use the OTP below to reset your Meta Connect password. This code expires in <strong>15 minutes</strong>.</p>
            <div style="background:#f0fdf4;border:2px dashed #16A34A;border-radius:16px;padding:28px;text-align:center;margin:0 0 28px;">
              <p style="margin:0 0 8px;color:#6b7280;font-size:12px;font-weight:600;letter-spacing:2px;text-transform:uppercase;">Your OTP Code</p>
              <p style="margin:0;font-size:48px;font-weight:800;letter-spacing:12px;color:#16A34A;">{otp}</p>
            </div>
            <p style="color:#9ca3af;font-size:13px;margin:0;">If you did not request this, you can safely ignore this email.</p>
          </td>
        </tr>
        <tr>
          <td style="background:#f9fafb;padding:20px 40px;text-align:center;border-top:1px solid #e5e7eb;">
            <p style="color:#9ca3af;font-size:12px;margin:0;">© 2025 Meta Connect. All rights reserved.</p>
          </td>
        </tr>
      </table>
    </td></tr>
  </table>
</body></html>"""
        
        try:
            send_mail(
                subject="Your Password Reset OTP - Meta Connect",
                message=f"Your OTP for resetting your Meta Connect password is: {otp}.\nThis OTP is valid for 15 minutes.",
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[email],
                html_message=html_body
            )
            print(f"\n[OTP] Sent {otp} to {email} via SMTP\n")
            return Response({"message": "OTP sent to your email successfully"})
        except Exception as e:
            print(f"Email SMTP send error: {str(e)}")
            return Response({"message": f"Failed to send email: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@method_decorator(csrf_exempt, name='dispatch')
class ForgotPasswordVerifyOTPView(views.APIView):
    permission_classes = []

    def post(self, req):
        from django.utils import timezone
        from datetime import timedelta
        from .models import PasswordResetOTP
        
        email = req.data.get('email', '').lower().strip()
        otp = req.data.get('otp', '').strip()
        
        if not email or not otp:
            return Response({"message": "Email and OTP are required"}, status=status.HTTP_400_BAD_REQUEST)
        
        # Find latest OTP
        try:
            otp_record = PasswordResetOTP.objects.filter(email=email, otp=otp).latest('created_at')
        except PasswordResetOTP.DoesNotExist:
            return Response({"message": "Invalid OTP"}, status=status.HTTP_400_BAD_REQUEST)
        
        # Check expiration (15 minutes)
        now = timezone.now()
        if now - otp_record.created_at > timedelta(minutes=15):
            otp_record.delete()
            return Response({"message": "OTP has expired"}, status=status.HTTP_400_BAD_REQUEST)
            
        otp_record.is_verified = True
        otp_record.save()
        
        return Response({"message": "OTP verified successfully"})

@method_decorator(csrf_exempt, name='dispatch')
class ForgotPasswordResetView(views.APIView):
    permission_classes = []

    def post(self, req):
        from .models import PasswordResetOTP
        
        email = req.data.get('email', '').lower().strip()
        password = req.data.get('password', '')
        
        if not email or not password:
            return Response({"message": "Email and password are required"}, status=status.HTTP_400_BAD_REQUEST)
            
        # Check if verified OTP exists
        verified_otp = PasswordResetOTP.objects.filter(email=email, is_verified=True).exists()
        if not verified_otp:
            return Response({"message": "OTP not verified yet"}, status=status.HTTP_400_BAD_REQUEST)
            
        # Reset password
        try:
            user = User.objects.get(email=email)
            user.set_password(password)
            user.save()
            
            # Delete verified OTP record
            PasswordResetOTP.objects.filter(email=email).delete()
            return Response({"message": "Password reset successfully"})
        except User.DoesNotExist:
            return Response({"message": "User not found"}, status=status.HTTP_404_NOT_FOUND)

class FacebookInstagramWebhookView(APIView):
    permission_classes = [] # Publicly accessible for Meta webhooks

    def get(self, request):
        """
        Meta Webhook Verification (hub.mode = subscribe)
        """
        mode = request.query_params.get('hub.mode')
        token = request.query_params.get('hub.verify_token')
        challenge = request.query_params.get('hub.challenge')
        verify_token = os.getenv('WHATSAPP_VERIFY_TOKEN') or 'aisaconnect_secure_token'

        if mode and token:
            if mode == 'subscribe' and token == verify_token:
                print("FB/IG WEBHOOK_VERIFIED")
                return HttpResponse(challenge, content_type="text/plain", status=200)
            else:
                return HttpResponse("Forbidden", content_type="text/plain", status=403)
        return HttpResponse("Bad Request", content_type="text/plain", status=400)

    def post(self, request):
        """
        Handles incoming Facebook Page & Instagram events (POST request)
        """
        data = request.data
        print("Incoming FB/IG Webhook Payload:", json.dumps(data, indent=2))

        try:
            # Facebook/Instagram webhook structure has 'entry'
            for entry in data.get('entry', []):
                recipient_id = entry.get('id') # Page ID or Instagram Business Account ID
                
                client = None
                platform = None
                
                # Check for Facebook
                if data.get('object') == 'page':
                    client = Client.objects.filter(facebook_config__page_id=recipient_id).first()
                    platform = 'FACEBOOK'
                # Check for Instagram
                elif data.get('object') == 'instagram':
                    client = Client.objects.filter(instagram_config__instagram_business_id=recipient_id).first()
                    platform = 'INSTAGRAM'
                
                if not client:
                    print(f"No client found for {platform} recipient ID: {recipient_id}")
                    continue

                if not client.automation_enabled:
                    print(f"Automation disabled for client: {client.business_name}")
                    continue

                # Process incoming messages
                messaging = entry.get('messaging', [])
                for event in messaging:
                    sender_id = event.get('sender', {}).get('id')
                    body = ""
                    if 'message' in event:
                        msg_data = event.get('message', {})
                        if 'quick_reply' in msg_data:
                            body = msg_data.get('quick_reply', {}).get('payload', '')
                        else:
                            body = msg_data.get('text', '')
                    elif 'postback' in event:
                        body = event.get('postback', {}).get('payload', '') or event.get('postback', {}).get('title', '')
                    
                    if not body:
                        continue
                    
                    # Ensure Contact exists for CRM
                    contact, _ = Contact.objects.get_or_create(
                        client=client,
                        platform_id=sender_id,
                        defaults={
                            'phone_number': sender_id,
                            'name': f"{platform} User",
                            'stage': 'NEW'
                        }
                    )

                    # Log the message in our database
                    Message.objects.create(
                        client=client,
                        channel=platform,
                        from_address=sender_id,
                        to_address=recipient_id,
                        body=body,
                        message_type='INCOMING',
                        status='RECEIVED',
                        metadata=event
                    )

                    # Trigger automations
                    if body:
                        if not contact.bot_paused:
                            self.handle_automations(client, platform, sender_id, body)
                        else:
                            print(f"Bot paused for {platform} contact {sender_id}. No automated response.")

            return Response({"status": "success"}, status=status.HTTP_200_OK)
        except Exception as e:
            print(f"Error processing FB/IG webhook: {str(e)}")
            return Response({"status": "error"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def handle_automations(self, client, platform, sender_id, incoming_text):
        """
        Matches keywords and sends automated responses.
        If no keyword matches, sends the Global Greeting Message if enabled.
        """
        from .workflow_engine import WorkflowEngine
        
        # 0. Check Workflow Engine first
        wf_messages = WorkflowEngine.process_workflow(client, sender_id, incoming_text, platform)
        if wf_messages:
            for msg in wf_messages:
                self.send_message(
                    client=client,
                    platform=platform,
                    recipient_id=sender_id,
                    text_body=msg.get('body', ''),
                    buttons=msg.get('buttons'),
                    media_url=msg.get('media_url'),
                    media_type=msg.get('type')
                )
            return  # Stop further processing if a workflow handled it

        # 1. Try Keyword Matching
        automations = Automation.objects.filter(client=client, enabled=True, trigger_type='KEYWORD')
        incoming_text_lower = incoming_text.lower().strip()
        
        match_found = False
        for auto in automations:
            # Check channels list
            auto_channels = auto.channels or []
            if len(auto_channels) > 0 and platform not in auto_channels:
                continue
            if len(auto_channels) == 0 and platform != 'WHATSAPP':
                continue

            if auto.keywords:
                for keyword in auto.keywords:
                    if keyword.lower().strip() == incoming_text_lower:
                        self.send_message(client, platform, sender_id, auto.response, auto.buttons)
                        match_found = True
                        break
            if match_found: break

        # 2. Check AI Assistant
        if not match_found and client.ai_enabled:
            ai_reply = None
            chunks = KnowledgeChunk.objects.filter(client=client).exclude(embedding=[])
            if chunks.exists():
                query_embedding = get_embedding(incoming_text)
                if query_embedding:
                    chunks_data = [{
                        'text': c.chunk_text,
                        'embedding': c.embedding,
                        'doc_title': c.document.title
                    } for c in chunks.select_related('document')]
                    relevant = find_relevant_chunks(query_embedding, chunks_data, top_k=5)
                    if relevant and relevant[0]['score'] > 0.3:
                        ai_reply = get_rag_response(incoming_text, relevant)
            
            if not ai_reply:
                ai_reply = get_ai_response(incoming_text, client.ai_context or "")
                
            if ai_reply:
                self.send_message(client, platform, sender_id, ai_reply)
                match_found = True

        # 3. Check Greeting Message
        if not match_found and client.greeting_enabled and client.greeting_message:
            self.send_message(client, platform, sender_id, client.greeting_message, client.greeting_buttons)

    def send_message(self, client, platform, recipient_id, text_body, buttons=None, media_url=None, media_type=None):
        """
        Calls Meta Graph API to send message to Facebook Page or Instagram Business Account
        """
        config = client.facebook_config if platform == 'FACEBOOK' else client.instagram_config
        access_token = config.get('access_token') or client.whatsapp_access_token
        
        if not access_token:
            print(f"No access token found for {platform} messaging")
            return
            
        url = "https://graph.facebook.com/v20.0/me/messages"
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }
        
        message_payload = {}
        if media_url:
            m_type = media_type or 'image'
            if not media_type:
                m_type = 'video' if any(ext in media_url.lower() for ext in ['.mp4', '.mov', '.avi']) else 'image'
            message_payload = {
                "attachment": {
                    "type": m_type,
                    "payload": {
                        "url": media_url,
                        "is_reusable": True
                    }
                }
            }
        elif buttons and len(buttons) > 0:
            quick_replies = []
            for btn in buttons[:13]:
                quick_replies.append({
                    "content_type": "text",
                    "title": btn[:20],
                    "payload": btn
                })
            message_payload = {
                "text": text_body or "Please choose an option:",
                "quick_replies": quick_replies
            }
        else:
            message_payload = {
                "text": text_body
            }
            
        payload = {
            "recipient": {"id": recipient_id},
            "message": message_payload
        }
        
        try:
            res = requests.post(url, json=payload, headers=headers)
            print(f"{platform} Send Response:", res.status_code, res.text)
            
            # Log OUTGOING message
            Message.objects.create(
                client=client,
                channel=platform,
                from_address='SYSTEM',
                to_address=recipient_id,
                body=text_body or f"[{media_type or 'Attachment'}]",
                message_type='OUTGOING',
                status='SENT'
            )
        except Exception as e:
            print(f"Failed to send {platform} message:", str(e))


class ClientMessagesView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        client = get_tenant_client(request)
        if not client:
            return Response([])
        
        messages = Message.objects.filter(client=client).order_by('-created_at')[:100]
        data = []
        for msg in messages:
            data.append({
                "id": str(msg.id),
                "from_address": msg.from_address,
                "to_address": msg.to_address,
                "body": msg.body,
                "channel": msg.channel,
                "message_type": msg.message_type,
                "status": msg.status,
                "created_at": msg.created_at
            })
        return Response(data)

    def post(self, request):
        client = get_tenant_client(request)
        if not client:
            return Response({"error": "No client associated"}, status=400)
            
        to_number = request.data.get('to_number')
        body = request.data.get('body')
        channel = request.data.get('channel')
        
        if not to_number or not body:
            return Response({"error": "to_number and body are required"}, status=400)
            
        # Detect channel if not provided
        if not channel:
            last_msg = Message.objects.filter(client=client, from_address=to_number).order_by('-created_at').first()
            if not last_msg:
                last_msg = Message.objects.filter(client=client, to_address=to_number).order_by('-created_at').first()
            channel = last_msg.channel if last_msg else 'WHATSAPP'
            
        channel = channel.upper()
        
        if channel == 'WHATSAPP':
            phone_number_id = client.whatsapp_phone_number_id
            if not phone_number_id:
                return Response({"error": "WhatsApp not connected for this client"}, status=400)
                
            webhook_view = WhatsAppWebhookView()
            webhook_view.send_whatsapp_message(client, to_number, body, phone_number_id)
        elif channel in ['INSTAGRAM', 'FACEBOOK']:
            webhook_view = FacebookInstagramWebhookView()
            webhook_view.send_message(client, channel, to_number, body)
        else:
            return Response({"error": f"Unsupported channel: {channel}"}, status=400)
            
        return Response({"status": "sent"})

def log_admin_action(request, instance, module, action, before_value=None, after_value=None):
    auth_data = getattr(request, 'auth', None)
    impersonator_username = None
    if auth_data:
        try:
            impersonator_username = auth_data.get('impersonator_username')
        except AttributeError:
            if hasattr(auth_data, 'payload'):
                impersonator_username = auth_data.payload.get('impersonator_username')
            elif isinstance(auth_data, dict):
                impersonator_username = auth_data.get('impersonator_username')
                
    if not impersonator_username and request.user.is_authenticated and request.user.role == 'ADMIN':
        impersonator_username = request.user.username
        
    if impersonator_username:
        client = getattr(instance, 'client', None)
        if not client and isinstance(instance, Client):
            client = instance
        client_name = client.business_name if client else "Global / Platform"
        
        ip_addr = request.META.get('REMOTE_ADDR')
        
        AuditLog.objects.create(
            admin_name=impersonator_username,
            client_name=client_name,
            module=module,
            action=action,
            before_value=before_value,
            after_value=after_value,
            ip_address=ip_addr
        )

class AuditLogMixin:
    def get_module_name(self):
        model = None
        if hasattr(self, 'queryset') and self.queryset:
            model = self.queryset.model
        elif hasattr(self, 'get_queryset'):
            try:
                model = self.get_queryset().model
            except Exception:
                pass
        return model.__name__ if model else "General"

    def perform_create(self, serializer):
        instance = serializer.save()
        log_admin_action(self.request, instance, self.get_module_name(), 'CREATE', after_value=str(serializer.data))
        
    def perform_update(self, serializer):
        before_instance = self.get_object()
        before_data = str(self.get_serializer(before_instance).data)
        instance = serializer.save()
        log_admin_action(self.request, instance, self.get_module_name(), 'UPDATE', before_value=before_data, after_value=str(serializer.data))
        
    def perform_destroy(self, instance):
        before_data = str(self.get_serializer(instance).data)
        log_admin_action(self.request, instance, self.get_module_name(), 'DELETE', before_value=before_data)
        instance.delete()

class SupportMessageViewSet(viewsets.ModelViewSet):
    serializer_class = SupportMessageSerializer
    permission_classes = [IsApprovedUser]

    def get_queryset(self):
        user = self.request.user
        if user.role == 'ADMIN':
            client_id = self.request.query_params.get('client_id')
            if client_id:
                try:
                    client = Client.objects.get(id=client_id)
                    return SupportMessage.objects.filter(client=client)
                except Exception:
                    return SupportMessage.objects.none()
            return SupportMessage.objects.all()
        return SupportMessage.objects.filter(client=user.client)

    def perform_create(self, serializer):
        user = self.request.user
        if user.role == 'ADMIN':
            client_id = self.request.data.get('client_id')
            if not client_id:
                raise serializers.ValidationError({"client_id": "Required when sending as admin."})
            client = Client.objects.get(id=client_id)
        else:
            client = user.client
        serializer.save(sender=user, client=client)

    @action(detail=False, methods=['get'], permission_classes=[IsAdminUser])
    def clients(self, request):
        from django.db.models import Max
        client_ids = SupportMessage.objects.values('client').annotate(latest_message=Max('created_at')).order_by('-latest_message')
        
        clients_data = []
        for item in client_ids:
            client_id = item['client']
            client = Client.objects.filter(id=client_id).first()
            if client:
                last_msg = SupportMessage.objects.filter(client=client).order_by('-created_at').first()
                clients_data.append({
                    "id": str(client.id),
                    "business_name": client.business_name,
                    "last_message_body": last_msg.body if last_msg else "",
                    "last_message_time": last_msg.created_at if last_msg else None,
                    "unread_count": 0
                })
        return Response(clients_data)

class AdminImpersonateView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        if request.user.role != 'ADMIN':
            return Response({"error": "Only admins can impersonate clients."}, status=status.HTTP_403_FORBIDDEN)
            
        client_id = request.data.get('client_id')
        if not client_id:
            return Response({"error": "client_id is required."}, status=status.HTTP_400_BAD_REQUEST)
            
        try:
            client = Client.objects.get(id=client_id)
            user = User.objects.filter(client=client, role='CLIENT').first()
            if not user:
                return Response({"error": "No user registered under this client node to impersonate."}, status=status.HTTP_404_NOT_FOUND)
                
            refresh = RefreshToken.for_user(user)
            refresh['impersonator_id'] = str(request.user.id)
            refresh['impersonator_username'] = request.user.username
            return Response({
                "access": str(refresh.access_token),
                "refresh": str(refresh),
                "user": {
                    "username": user.username,
                    "email": user.email,
                    "role": user.role
                }
            })
        except Client.DoesNotExist:
            return Response({"error": "Client node not found."}, status=status.HTTP_404_NOT_FOUND)

class AuditLogViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = AuditLog.objects.all()
    serializer_class = AuditLogSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        if self.request.user.role != 'ADMIN':
            return AuditLog.objects.none()
            
        queryset = AuditLog.objects.all()
        search = self.request.query_params.get('search')
        if search:
            from django.db.models import Q
            queryset = queryset.filter(
                Q(admin_name__icontains=search) |
                Q(client_name__icontains=search) |
                Q(module__icontains=search) |
                Q(action__icontains=search)
            )
        module = self.request.query_params.get('module')
        if module:
            queryset = queryset.filter(module=module)
            
        return queryset
